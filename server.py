"""
AI 법률 문서 자동화 서비스 — Flask 백엔드

역할:
  · 브라우저는 정적인 HTML/CSS/JS 만 다룬다 (API 키 노출 X)
  · 모든 Claude API 호출과 .docx 생성은 이 서버가 담당한다

엔드포인트:
  GET  /                      → templates/index.html 렌더링
  POST /api/generate          → JSON 입력 → Claude 초안 생성 → JSON 반환
  POST /api/revise            → 초안 + 수정 요청 → Claude 수정 → JSON 반환
  POST /api/download_docx     → 편집된 텍스트 → .docx 파일 응답

실행:
    pip install flask anthropic python-docx python-dotenv
    python server.py
브라우저:
    http://localhost:5001
"""

import io
import os
import re
from datetime import date, datetime

from dotenv import load_dotenv
load_dotenv(override=True)

from flask import Flask, jsonify, render_template, request, send_file, send_from_directory
import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from doc_templates import DOC_TEMPLATES
from payment_routes import (
    payment_bp,
    check_document_access,
    record_generate_use,
    check_and_increment_revision,
    get_trial_status,
    _normalize_doc_type,
    DOC_PRICES,
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(
    __name__,
    template_folder=os.path.join(BASE_DIR, "templates"),
    static_folder=os.path.join(BASE_DIR, "static"),
    static_url_path="/static",
)

# 건별 결제 API 라우트 등록 (PRD 건별결제 v1.0)
app.register_blueprint(payment_bp)


def _gate_user_id():
    """결제 권한 게이트용 사용자 식별 (Phase 1: 헤더/바디, Phase 2: Supabase JWT)."""
    uid = request.headers.get("X-User-Id")
    if not uid:
        body = request.get_json(silent=True) or {}
        uid = body.get("user_id")
    return uid or "anonymous"


# ── Supabase 클라이언트 (documents 테이블 연동, 설계 PRD §2~3) ─────────
#
# [수동 작업 안내] 아래 DDL 을 Supabase 대시보드 SQL Editor 에서 1회 실행해
# documents 테이블·RLS·트리거를 생성해야 합니다. (서버는 테이블을 만들지 않음)
#
#   create table public.documents (
#     id            uuid primary key default gen_random_uuid(),
#     user_id       uuid not null references auth.users(id) on delete cascade,
#     doc_type      text not null check (doc_type in ('notice','brief','rebuttal')),
#     title         text not null default '(제목 없음)',
#     status        text not null default 'draft'
#                     check (status in ('draft','generated','in_review','saved','delivered','deleted')),
#     current_step  smallint not null default 1 check (current_step between 1 and 3),
#     input_data    jsonb,
#     draft_text    text,
#     revision_count smallint not null default 0,
#     created_at    timestamptz not null default now(),
#     updated_at    timestamptz not null default now()
#   );
#   alter table public.documents enable row level security;
#   create policy "own documents" on public.documents
#     using (auth.uid() = user_id) with check (auth.uid() = user_id);
#   create or replace function update_updated_at()
#   returns trigger language plpgsql as $$
#   begin new.updated_at = now(); return new; end; $$;
#   create trigger documents_updated_at
#     before update on public.documents
#     for each row execute function update_updated_at();
#
# [인증 경계 — Phase 1]
#   서버는 JWT 가 아닌 X-User-Id 헤더로 사용자를 식별한다(_gate_user_id).
#   anon key 클라이언트로는 RLS auth.uid() 를 만족시킬 수 없으므로, 모든 쿼리에
#   명시적으로 .eq("user_id", user_id) 필터를 걸어 1차 경계로 삼는다.
#   RLS 는 Phase 2(사용자 JWT 전달 또는 service-role 키)에서 2차 방어로 활성화한다.

_SUPABASE_CLIENT = None


def _get_supabase():
    """documents 테이블용 Supabase 클라이언트(싱글턴). 미설정 시 None 반환."""
    global _SUPABASE_CLIENT
    if _SUPABASE_CLIENT is not None:
        return _SUPABASE_CLIENT
    url = (os.environ.get("SUPABASE_URL") or "").strip()
    key = (os.environ.get("SUPABASE_SERVICE_KEY")
           or os.environ.get("SUPABASE_ANON_KEY") or "").strip()
    if not url or not key:
        return None
    # .env 의 URL 에 스킴이 없을 수 있으므로 보정
    if not url.startswith("http"):
        url = "https://" + url
    try:
        from supabase import create_client
        _SUPABASE_CLIENT = create_client(url, key)
    except Exception:
        return None
    return _SUPABASE_CLIENT


# documents 목록 응답에서 제외할 무거운 필드 (설계 PRD §3.2)
_DOC_LIST_COLUMNS = (
    "id,doc_type,title,status,current_step,revision_count,created_at,updated_at"
)
_VALID_DOC_TYPES = {"notice", "brief", "rebuttal"}
_VALID_DOC_STATUS = {"draft", "generated", "in_review", "saved", "delivered", "deleted"}
_PATCHABLE_FIELDS = {
    "title", "status", "current_step", "input_data", "draft_text", "revision_count",
}


# ── 한글(East Asian) 폰트 적용 헬퍼 ──────────────────────────
def _set_korean_font(run, font_name: str) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _format_timeline(events):
    rows = []
    for ev in events or []:
        t = (ev.get("time") or "").strip()
        c = (ev.get("content") or "").strip()
        if t or c:
            rows.append(f"  · [{t or '시점 미상'}] {c or '(내용 미입력)'}")
    return "\n".join(rows) if rows else "(미입력)"


def build_prompt(doc_type, sender, receiver, case_info,
                 timeline_events, facts, request_text):
    today = date.today().strftime("%Y년 %m월 %d일")
    timeline_block = _format_timeline(timeline_events)
    has_timeline = timeline_block != "(미입력)"
    timeline_instruction = (
        "\n6. [시간순별 사건 경위] 의 각 항목을 본문의 사실관계/사건 경위 섹션에\n"
        "   반드시 시간 순서대로 통합·서술할 것. 시각([…])을 그대로 인용하고,\n"
        "   사건내용은 격식 있는 문장으로 풀어 쓰며, '먼저', '이후', '그 다음',\n"
        "   '마지막으로' 등 접속어를 자연스럽게 사용할 것.\n"
        if has_timeline else ""
    )
    return f"""당신은 한국 법률 실무에 능통한 전문 법률 보조원입니다.
아래 [의뢰인 정보]를 바탕으로 「{doc_type}」 초안을 작성해 주세요.

반드시 아래 [문서 양식 가이드]의 구조와 정형 문구를 그대로 지키고,
실제 우체국 또는 법원에 제출 가능한 수준의 격식 있는 한국어로 작성하십시오.

[문서 양식 가이드]
{DOC_TEMPLATES[doc_type]["format_guide"]}

[의뢰인 정보]
- 작성일자(오늘): {today}
- 발신인/작성자: {sender or "(미입력)"}
- 수신인: {receiver or "(미입력)"}
- 사건 표시: {case_info or "(해당 없음)"}
- 시간순별 사건 경위 (시간 → 사건내용):
{timeline_block}
- 사건 경위 및 사실관계 (자유 서술):
{facts or "(미입력)"}
- 요구사항/주장 결론:
{request_text or "(미입력)"}

[작성 지시]
1. 양식 가이드의 모든 섹션을 빠짐없이 포함할 것.
2. 각 섹션 제목(예: 발 신 인, - 아 래 -, - 끝 -, 다 음, 증 명 방 법,
   첨 부 서 류, 항 소 취 지, 항 소 이 유 등)은 그대로 한 줄로 출력할 것.
3. 본문은 자연스러운 한국어 격식체로 작성하고, 추측 표현은 최소화할 것.
4. 마크다운 기호(#, *, -, > 등)나 코드블록을 사용하지 말 것.
5. 문서의 마지막 줄에 다음 안내 문구를 별도 단락으로 추가할 것:
   「본 문서는 법률적 참고용 초안이며, 실제 제출 전 반드시 전문가와 상담하십시오.」
6. 본문의 사실관계나 주장 부분에서 관련 증거가 명확한 경우 괄호 안에 증거번호를
   자연스럽게 인용하십시오. 단, 문서 하단의 입증방법 목록은 서버에서 자동으로
   추가하므로 직접 작성하지 마십시오.{timeline_instruction}
"""


CENTER_KEYWORDS = (
    "내 용 증 명", "내  용  증  명",
    "소 견 서", "소  견  서",
    "준 비 서 면", "준  비  서  면",
    "항 소 이 유 서", "항   소   이   유   서",
    "- 아  래 -", "- 아 래 -", "- 아래 -", "- 끝 -",
    "다 음", "다  음", "다        음",
    "항 소 취 지", "항   소   취   지",
    "항 소 이 유", "항   소   이   유",
    "청 구 취 지", "청   구   취   지",
    "청 구 원 인", "청   구   원   인",
    "증 명 방 법", "증   명   방   법",
    "입 증 방 법", "입   증   방   법",
    "소 명 자 료", "소  명  자  료",
    "첨 부 서 류", "첨   부   서   류",
)
RIGHT_KEYWORDS = ("(인)", "대표이사")


def make_docx(text: str, title: str, watermark_text: str = None) -> bytes:
    BODY_FONT = "바탕체"
    TITLE_FONT = "맑은 고딕"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    spr = style.element.get_or_add_rPr()
    sf = spr.find(qn("w:rFonts"))
    if sf is None:
        sf = OxmlElement("w:rFonts")
        spr.insert(0, sf)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        sf.set(qn(attr), BODY_FONT)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for raw_line in text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        stripped = line.strip()
        is_title_line = (stripped in CENTER_KEYWORDS
                         or stripped == title
                         or stripped.endswith("귀중"))
        if stripped in CENTER_KEYWORDS or stripped == title:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True
            if (stripped == title
                or any(t in stripped for t in
                       ("내 용 증 명", "소 견 서", "준 비 서 면", "항 소 이 유 서"))):
                run.font.size = Pt(18)
        elif any(stripped.startswith(k) or stripped.endswith(k) for k in RIGHT_KEYWORDS):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif re.match(r"^\d{4}\s*[\.년]\s*\d{1,2}\s*[\.월]\s*\d{1,2}", stripped):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif stripped.endswith("귀중"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_korean_font(run, TITLE_FONT if is_title_line else BODY_FONT)

    if watermark_text:
        doc.add_paragraph("")
        wm_p = doc.add_paragraph()
        wm_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        wm_run = wm_p.add_run(f"[ {watermark_text} ]")
        wm_run.bold = True
        wm_run.font.size = Pt(9)
        _set_korean_font(wm_run, BODY_FONT)
        # 구분선 역할 상단 테두리
        pPr = wm_p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "4")
        top.set(qn("w:space"), "4")
        top.set(qn("w:color"), "AAAAAA")
        pBdr.append(top)
        pPr.append(pBdr)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── 증거자료 저장·명명 (PRD: 증거 자료 명기 및 저장) ──────────────
EVIDENCE_DIR = os.path.join(BASE_DIR, "증거파일")

# 문서 종류별 증거번호 접두 (준비서면→갑, 반박문→을, 내용증명→첨부)
EVIDENCE_PREFIX = {"brief": "갑", "rebuttal": "을", "notice": "첨부"}


def _evidence_label_no(doc_type: str, seq: int):
    """증거번호 표기와 파일명 토큰을 함께 반환. (표기, 파일토큰)"""
    prefix = EVIDENCE_PREFIX.get(doc_type, "갑")
    if prefix == "첨부":
        return (f"첨부 제{seq}호", f"첨부{seq}")
    return (f"{prefix} 제{seq}호증", f"{prefix}제{seq}호증")


def _sanitize_part(s: str, maxlen: int = 24) -> str:
    """파일명 토큰에서 공백·경로·금지문자 제거 (path traversal 방지 포함)."""
    s = (s or "").strip()
    # 경로 구분자/상위경로/윈도우 금지문자/제어문자/공백 제거
    s = re.sub(r"[\\/:*?\"<>|\x00-\x1f]", "", s)
    s = s.replace("..", "").replace("。", "")  # 상위경로·전각마침표 방지
    s = re.sub(r"\s+", "", s)
    return s[:maxlen] or "자료"


# 증거번호 토큰(갑제1호증 / 을제2호증 / 증1 / 첨부1 등) 판별 — 자료명 중복 제거용
_EVIDENCE_NO_RE = re.compile(r"^(갑|을|증|첨부)제?\d+호?증?$")


def _clean_name_token(label: str, maxlen: int = 24) -> str:
    """자료명 토큰 정리. 이미 명명된 파일 재업로드 시 증거번호·'날짜미상' 토큰을
    제거해 '첨부1_첨부1_…_날짜미상' 같은 중복을 방지한다."""
    base = os.path.splitext(label or "")[0]
    kept = []
    for tok in re.split(r"[_\s]+", base):
        t = _sanitize_part(tok, maxlen)
        if not t or t == "자료" or t == "날짜미상":
            continue
        if _EVIDENCE_NO_RE.match(t):
            continue
        kept.append(t)
    name = "".join(kept)[:maxlen]
    return name or "자료"


def _unique_path(directory: str, filename: str) -> str:
    """동일 파일명 충돌 시 _2, _3… 부여 (기존 파일 덮어쓰기 방지)."""
    base, ext = os.path.splitext(filename)
    candidate, n = filename, 2
    while os.path.exists(os.path.join(directory, candidate)):
        candidate = f"{base}_{n}{ext}"
        n += 1
    return candidate


def format_evidence_section(evidence_list):
    """업로드된 증거 목록을 문서 하단 「입 증 방 법」 섹션으로 결정적 생성.
    (AI 프롬프트에 맡기지 않고 서버에서 고정 양식으로 붙여 문구·순서를 보장)"""
    if not evidence_list:
        return ""

    lines = ["", "", "입 증 방 법", ""]
    for idx, ev in enumerate(evidence_list, start=1):
        evidence_no = (
            ev.get("evidenceNo")
            or ev.get("evidence_no")
            or f"증 제{idx}호증"
        )
        title = (
            ev.get("savedFilename")
            or ev.get("saved_filename")
            or ev.get("originalName")
            or ev.get("original_name")
            or "증거자료"
        )
        summary = ev.get("summary") or ev.get("description") or ""
        lines.append(f"{idx}. {evidence_no}  {title}")
        if summary:
            lines.append(f"   - 입증취지: {summary}")
        lines.append("")

    return "\n".join(lines).rstrip()


def _has_evidence_section(draft_text):
    """초안에 '입증방법' 섹션 헤더 줄이 이미 있는지 검사.
    본문 상용구(예: 첨부서류의 '위 입증방법 각 1통')는 오탐하지 않도록,
    공백 제거 후 한 줄이 정확히 '입증방법'인 헤더 줄만 인정한다."""
    for line in draft_text.split("\n"):
        if line.replace(" ", "").strip() == "입증방법":
            return True
    return False


def append_evidence_section(draft_text, evidence_list):
    """초안 본문에 증거목록을 1회만 안전하게 추가 (이미 섹션이 있으면 중복 방지)."""
    if evidence_list and not _has_evidence_section(draft_text):
        return draft_text + format_evidence_section(evidence_list)
    return draft_text


def _strip_preview_only_notes(text: str) -> str:
    """미리보기 전용 고지문구만 다운로드 본문에서 제거 (PRD: 다운로드 문서 형식).
    과도 삭제 방지를 위해 '참고용 초안' + (상담/전문가/검토)가 함께 있는 줄만 제거한다."""
    out = []
    for line in text.split("\n"):
        s = line.strip().strip("「」 ")
        if "참고용 초안" in s and ("상담" in s or "전문가" in s or "검토" in s):
            continue
        out.append(line)
    while out and not out[-1].strip():
        out.pop()
    return "\n".join(out)


# ── 라우트 ────────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template(
        "index.html",
        supabase_url=os.environ.get("SUPABASE_URL", ""),
        supabase_anon_key=os.environ.get("SUPABASE_ANON_KEY", ""),
        # 결제 클라이언트 설정 (공개값만 — SECRET/WEBHOOK 키는 절대 전달 금지, NFR-PAY-01)
        pg_test_mode=(os.environ.get("PG_TEST_MODE", "true").lower() == "true"),
        pg_client_key=os.environ.get("PG_CLIENT_KEY", ""),
        pg_provider=os.environ.get("PG_PROVIDER", "mock"),
    )


@app.route("/images/<path:filename>")
def serve_image(filename):
    img_dir = os.path.join(BASE_DIR, "images")
    return send_from_directory(img_dir, filename)


@app.route("/api/generate", methods=["POST"])
def api_generate():
    data = request.get_json(force=True)
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "서버에 Anthropic API 키가 설정되지 않았습니다. .env 파일을 확인해 주세요."}), 500

    doc_type = data.get("doc_type")
    # 영문 키(notice/brief/rebuttal)가 그대로 오는 경우 한글로 변환 (방어적 처리)
    _EN_TO_KOR = {"notice": "내용증명", "brief": "준비서면", "rebuttal": "상대방 반박문"}
    doc_type = _EN_TO_KOR.get(doc_type, doc_type)
    if doc_type not in DOC_TEMPLATES:
        return jsonify({"error": f"알 수 없는 문서 종류: {doc_type}"}), 400
    if not (data.get("facts") or "").strip():
        return jsonify({"error": "사건 경위(facts) 는 필수입니다."}), 400

    # FR-30: 문서 접근 권한 게이트 (결제 대상 3종 한정). 권한 없으면 403 PAYMENT_REQUIRED.
    user_id = _gate_user_id()
    doc_type_en = _normalize_doc_type(doc_type)
    if doc_type_en in DOC_PRICES:
        access = check_document_access(user_id, doc_type_en)
        if not access["allowed"]:
            return jsonify({"error": "PAYMENT_REQUIRED",
                            "message": "결제가 필요한 문서입니다."}), 403

    prompt = build_prompt(
        doc_type=doc_type,
        sender=data.get("sender", ""),
        receiver=data.get("receiver", ""),
        case_info=data.get("case_info", ""),
        timeline_events=data.get("timeline_events", []),
        facts=data.get("facts", ""),
        request_text=data.get("request", ""),
    )
    evidence_list = data.get("evidence_list", []) or []
    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=8000,
            messages=[{"role": "user", "content": prompt}],
        )
        draft = resp.content[0].text.strip()
        # 증거목록은 모델이 아닌 서버에서 결정적으로 1회 append (중복 방지)
        draft = append_evidence_section(draft, evidence_list)
        # FR-30/FR-28: 생성 성공 시 사용 횟수 기록 (환불 가능 여부 판단 기준)
        if doc_type_en in DOC_PRICES:
            record_generate_use(user_id, doc_type_en)
        return jsonify({"draft": draft})
    except Exception as e:
        return jsonify({"error": f"Claude 호출 실패: {e}"}), 500


@app.route("/api/revise", methods=["POST"])
def api_revise():
    data = request.get_json(force=True)
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "서버에 Anthropic API 키가 설정되지 않았습니다. .env 파일을 확인해 주세요."}), 500

    draft = (data.get("draft") or "").strip()
    revision_request = (data.get("revision_request") or "").strip()

    if not draft:
        return jsonify({"error": "수정할 초안이 없습니다."}), 400
    if not revision_request:
        return jsonify({"error": "수정 요청 내용을 입력해주세요."}), 400

    # FR-30: 수정 횟수 게이트 (결제 대상 3종 한정)
    doc_type_en = _normalize_doc_type(data.get("doc_type") or data.get("docType"))
    if doc_type_en and doc_type_en in DOC_PRICES:
        user_id = _gate_user_id()
        rev = check_and_increment_revision(user_id, doc_type_en)
        if not rev["allowed"]:
            return jsonify({
                "error": "REVISION_LIMIT_EXCEEDED",
                "used": rev["used"],
                "limit": rev["limit"],
                "message": f"대화형 수정 횟수({rev['limit']}회)를 모두 사용하셨습니다.",
            }), 403

    prompt = f"""다음은 한국 법률 문서 초안입니다:

{draft}

사용자가 아래와 같이 수정을 요청했습니다:
{revision_request}

위 수정 요청을 반영하여 법률 문서 초안을 수정해주세요.
- 수정 요청에 언급된 내용만 변경하고, 나머지는 원본을 그대로 유지하세요.
- 격식 있는 한국어 법률 문서 양식을 유지하세요.
- 마크다운 기호(#, *, -, > 등)를 사용하지 마세요.
- 수정된 문서 전체를 반환해주세요."""

    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=8000,
            messages=[{"role": "user", "content": prompt}],
        )
        return jsonify({"draft": resp.content[0].text.strip()})
    except Exception as e:
        return jsonify({"error": f"Claude 호출 실패: {e}"}), 500


@app.route("/api/suggest_strategies", methods=["POST"])
def suggest_strategies():
    """초안 설득력 강화 전략 2종 제안 (FR-24 내편 전략 제안)"""
    data = request.get_json(force=True)
    draft = (data.get("draft") or "").strip()
    if not draft:
        return jsonify({"error": "초안이 없습니다"}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "서버에 Anthropic API 키가 설정되지 않았습니다. .env 파일을 확인해 주세요."}), 500

    prompt = f"""다음 법률 문서 초안을 분석하여 설득력을 높일 수 있는 전략 2가지를 제안해주세요.

초안:
{draft}

각 전략은 아래 JSON 형식으로 반환하세요:
{{
  "strategies": [
    {{"title": "전략 제목 (15자 이내)", "description": "전략 설명 (50자 이내, 구체적 방향)"}},
    {{"title": "전략 제목 (15자 이내)", "description": "전략 설명 (50자 이내, 구체적 방향)"}}
  ]
}}

JSON만 반환하고 다른 텍스트는 포함하지 마세요."""

    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}],
        )
        import json as _json, re as _re
        raw = resp.content[0].text.strip()
        # Claude가 ```json ... ``` 블록으로 감싸 반환하는 경우 제거
        raw = _re.sub(r'^```(?:json)?\s*', '', raw)
        raw = _re.sub(r'\s*```$', '', raw).strip()
        result = _json.loads(raw)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": f"전략 생성 실패: {e}"}), 500


@app.route("/api/download_docx", methods=["POST"])
def api_download_docx():
    data = request.get_json(force=True)
    text = data.get("text", "")
    title = data.get("title", "법률문서")
    if not text.strip():
        return jsonify({"error": "본문이 비어 있습니다."}), 400
    # 다운로드 본문에서만 미리보기 전용 고지문구 제거 (미리보기 화면은 그대로 유지)
    text = _strip_preview_only_notes(text)

    # M-08 / FR-30: 무료 체험 문서 워터마크 서버 사이드 삽입
    watermark_text = None
    doc_type_en = _normalize_doc_type(data.get("doc_type") or data.get("docType"))
    if doc_type_en and doc_type_en in DOC_PRICES:
        user_id = _gate_user_id()
        access = check_document_access(user_id, doc_type_en)
        if access.get("access_type") == "free_trial":
            watermark_text = (
                "본 문서는 내편문서 무료 체험 초안입니다. "
                "실제 제출 전 결제 후 워터마크 없는 문서를 발급받으세요."
            )

    blob = make_docx(text, title, watermark_text=watermark_text)
    # 다운로드 파일명: 문서종류명_날짜시간.docx (예: 내용증명_20260614_1530.docx)
    filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return send_file(
        io.BytesIO(blob),
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/api/upload_evidence", methods=["POST"])
def api_upload_evidence():
    """증거 파일 업로드 + AI 날짜·자료명·내용 추출 + 명명규칙 저장 (FR-17 / PRD 증거 명기)

    multipart/form-data:
      file    : 업로드 파일
      docType : notice|brief|rebuttal (증거번호 접두 결정)
      seq     : 증거 순번(1-based)
    AI 추출 실패(크레딧 부족 등) 시에도 fallback 값으로 명명·저장은 정상 동작한다.
    """
    if "file" not in request.files:
        return jsonify({"error": "파일이 없습니다."}), 400

    f = request.files["file"]
    filename = f.filename or "unknown"
    doc_type = request.form.get("docType", "brief")
    try:
        seq = int(request.form.get("seq", "1"))
    except (TypeError, ValueError):
        seq = 1

    raw = f.read()
    ext = os.path.splitext(filename)[1] or ""

    # ── AI 추출 (날짜·자료명·요약) — 실패해도 폴백으로 진행 ──
    date_val, label_val, summary_val = None, "", ""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if api_key:
        try:
            import base64, json as _json
            b64 = base64.standard_b64encode(raw).decode()
            mime = f.content_type or "application/octet-stream"
            instr = (
                "이 문서에서 (1) 가장 중요한 날짜(YYYY-MM-DD), "
                "(2) 자료 종류를 5~10자 한글 명사로(label, 예: 카카오톡대화내역/계약서/세금계산서), "
                "(3) 핵심 내용 한 줄(summary)을 추출하세요. "
                "JSON으로만 응답: {\"date\": \"YYYY-MM-DD\", \"label\": \"...\", \"summary\": \"...\"}"
            )
            if mime.startswith("image/"):
                content = [
                    {"type": "image", "source": {"type": "base64", "media_type": mime, "data": b64}},
                    {"type": "text", "text": instr},
                ]
            else:
                content = [{"type": "text", "text": f"파일명: {filename}\n{instr}"}]
            client = anthropic.Anthropic(api_key=api_key)
            resp = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=300,
                messages=[{"role": "user", "content": content}],
            )
            raw_text = resp.content[0].text.strip()
            try:
                parsed = _json.loads(raw_text)
            except Exception:
                parsed = {}
            date_val = parsed.get("date")
            label_val = parsed.get("label") or ""
            summary_val = parsed.get("summary") or ""
        except Exception:
            pass  # AI 실패 → 아래 폴백 사용

    # 폴백: 자료명은 원본 파일명(확장자 제외)
    if not label_val:
        label_val = os.path.splitext(filename)[0]

    # ── 명명 규칙 적용 (가이드에서 날짜 제외): [증거번호]_[자료명]_[간단설명].ext ──
    evidence_no, no_token = _evidence_label_no(doc_type, seq)
    name_token = _clean_name_token(label_val)
    desc_token = _clean_name_token(summary_val) if summary_val else ""
    if desc_token in ("자료", name_token):
        desc_token = ""
    parts = [no_token, name_token] + ([desc_token] if desc_token else [])
    saved_filename = "_".join(parts) + ext

    # ── 저장 (폴더 없으면 생성, 중복 시 _2… 부여) ──
    os.makedirs(EVIDENCE_DIR, exist_ok=True)
    saved_filename = _unique_path(EVIDENCE_DIR, saved_filename)
    with open(os.path.join(EVIDENCE_DIR, saved_filename), "wb") as out_f:
        out_f.write(raw)

    from urllib.parse import quote
    return jsonify({
        "filename": filename,
        "evidenceNo": evidence_no,
        "savedFilename": saved_filename,
        "extractedDate": date_val,
        "label": label_val,
        "summary": summary_val,
        "downloadUrl": "/api/download_evidence/" + quote(saved_filename),
    })


@app.route("/api/download_evidence/<path:filename>")
def api_download_evidence(filename):
    """저장된 증거 파일 다운로드. send_from_directory가 경로 이탈(path traversal)을 차단한다."""
    if not os.path.isdir(EVIDENCE_DIR):
        return jsonify({"error": "저장된 증거 파일이 없습니다."}), 404
    return send_from_directory(EVIDENCE_DIR, filename, as_attachment=True)


@app.route("/api/analyze_opponent", methods=["POST"])
def api_analyze_opponent():
    """상대방 문서 분석 + 반박 목록 생성 (FR-19)"""
    data = request.get_json(force=True)
    doc_kind  = data.get("docKind", "문서")
    doc_text  = data.get("docText", "")   # 텍스트로 붙여넣은 경우
    my_facts  = data.get("myFacts", "")   # 내 사건 경위 (컨텍스트용)

    if not doc_text.strip():
        return jsonify({"error": "분석할 문서 내용이 없습니다."}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API 키 미설정"}), 500

    prompt = f"""당신은 법률 전문가입니다.
아래 상대방 {doc_kind} 내용을 분석하여 주요 주장을 추출하고, 각 주장에 대한 반박 초안을 작성해 주세요.

[상대방 문서]
{doc_text}

[내 사건 경위 (참고)]
{my_facts}

다음 JSON 형식으로만 응답하세요:
{{
  "claims": [
    {{
      "claim": "상대방 주장 요약",
      "rebuttal": "반박 초안",
      "riskLevel": "normal 또는 caution (인정 위험이 있으면 caution)"
    }}
  ],
  "summary": "전체 반박 주장 통합 초안 (2-3문단)"
}}"""

    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}],
        )
        import json as _json
        raw_text = resp.content[0].text.strip()
        try:
            result = _json.loads(raw_text)
        except Exception:
            result = {"claims": [], "summary": raw_text}
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": f"분석 실패: {e}"}), 500


@app.route("/api/user/trial_status", methods=["GET"])
def api_trial_status():
    """PRD §8.2: 무료 체험 잔여 상태 조회 (FreeTrialBanner 연동)."""
    user_id = _gate_user_id()
    return jsonify(get_trial_status(user_id))


@app.route("/api/purchase", methods=["POST"])
def api_purchase():
    """건별 구매 기록 생성 (FR-21) — 실제 결제 게이트웨이 연동 전 목업"""
    data = request.get_json(force=True)
    doc_type   = data.get("docType", "")
    price      = data.get("price", 0)
    method     = data.get("method", "card")

    if not doc_type:
        return jsonify({"error": "문서 종류 누락"}), 400

    # 실제 구현 시: 결제 게이트웨이 호출 후 DB에 구매 기록 저장
    return jsonify({
        "success": True,
        "purchaseId": f"PUR-{date.today().strftime('%Y%m%d')}-MOCK",
        "docType": doc_type,
        "price": price,
        "method": method,
        "message": "구매가 완료되었습니다.",
    })


# ── 나의 문서 / 진행현황 DB 연동 API (설계 PRD §3, FR-10·FR-16) ─────
#
# 5종: GET /api/documents, GET /api/documents/stats,
#      POST /api/documents, PATCH /api/documents/<id>, DELETE /api/documents/<id>
# 공통: _gate_user_id() 인증 게이트 + .eq("user_id", user_id) 소유권 필터.

def _require_login():
    """로그인 사용자 id 반환. 미인증('anonymous')이면 None."""
    uid = _gate_user_id()
    return None if uid == "anonymous" else uid


@app.route("/api/documents", methods=["GET"])
def api_documents_list():
    """문서 목록 (필터·검색·정렬·페이지). draft_text/input_data 제외 (§3.2)."""
    user_id = _require_login()
    if not user_id:
        return jsonify({"error": "UNAUTHORIZED", "message": "로그인이 필요합니다."}), 401
    sb = _get_supabase()
    if sb is None:
        return jsonify({"error": "SUPABASE_UNAVAILABLE",
                        "message": "문서 저장소에 연결할 수 없습니다."}), 503

    try:
        page = max(1, int(request.args.get("page", 1)))
    except (TypeError, ValueError):
        page = 1
    try:
        per_page = min(100, max(1, int(request.args.get("per_page", 20))))
    except (TypeError, ValueError):
        per_page = 20

    doc_type = request.args.get("doc_type")
    status = request.args.get("status")
    q = (request.args.get("q") or "").strip()
    sort = request.args.get("sort", "updated_at:desc")

    sort_field, _, sort_dir = sort.partition(":")
    if sort_field not in ("created_at", "updated_at", "title"):
        sort_field = "updated_at"
    desc = (sort_dir or "desc").lower() != "asc"

    try:
        query = (
            sb.table("documents")
            .select(_DOC_LIST_COLUMNS, count="exact")
            .eq("user_id", user_id)
        )
        # status 미지정 시 삭제건 제외, 지정 시 해당 상태만
        if status and status in _VALID_DOC_STATUS:
            query = query.eq("status", status)
        else:
            query = query.neq("status", "deleted")
        if doc_type and doc_type in _VALID_DOC_TYPES:
            query = query.eq("doc_type", doc_type)
        if q:
            query = query.ilike("title", f"%{q}%")

        query = query.order(sort_field, desc=desc)
        start = (page - 1) * per_page
        query = query.range(start, start + per_page - 1)

        res = query.execute()
        total = res.count if res.count is not None else len(res.data or [])
        return jsonify({
            "items": res.data or [],
            "total": total,
            "page": page,
            "per_page": per_page,
        })
    except Exception as e:
        return jsonify({"error": f"문서 목록 조회 실패: {e}"}), 500


@app.route("/api/documents/stats", methods=["GET"])
def api_documents_stats():
    """통계 4종: 이번달 생성 / 저장완료 / 작성중 / 무료체험 잔여 (§3.3)."""
    user_id = _require_login()
    if not user_id:
        return jsonify({"error": "UNAUTHORIZED", "message": "로그인이 필요합니다."}), 401
    sb = _get_supabase()
    if sb is None:
        return jsonify({"error": "SUPABASE_UNAVAILABLE",
                        "message": "문서 저장소에 연결할 수 없습니다."}), 503

    def _count(builder):
        res = builder.execute()
        return res.count if res.count is not None else len(res.data or [])

    try:
        month_start = date.today().replace(day=1).isoformat()
        this_month = _count(
            sb.table("documents").select("id", count="exact")
            .eq("user_id", user_id).neq("status", "deleted")
            .gte("created_at", month_start)
        )
        saved = _count(
            sb.table("documents").select("id", count="exact")
            .eq("user_id", user_id).in_("status", ["saved", "delivered"])
        )
        in_progress = _count(
            sb.table("documents").select("id", count="exact")
            .eq("user_id", user_id).in_("status", ["draft", "generated", "in_review"])
        )
    except Exception as e:
        return jsonify({"error": f"통계 조회 실패: {e}"}), 500

    # 무료체험 잔여는 기존 결제 모듈 단일 출처 재활용 (§3.3, 내용증명 1건 한정)
    trial = get_trial_status(user_id)
    free_trial_remaining = 0 if trial.get("free_trial_used") else 1

    return jsonify({
        "this_month": this_month,
        "saved": saved,
        "in_progress": in_progress,
        "free_trial_remaining": free_trial_remaining,
    })


@app.route("/api/documents", methods=["POST"])
def api_documents_create():
    """문서 신규 생성. body: { doc_type } (§3.4)."""
    user_id = _require_login()
    if not user_id:
        return jsonify({"error": "UNAUTHORIZED", "message": "로그인이 필요합니다."}), 401
    sb = _get_supabase()
    if sb is None:
        return jsonify({"error": "SUPABASE_UNAVAILABLE",
                        "message": "문서 저장소에 연결할 수 없습니다."}), 503

    data = request.get_json(silent=True) or {}
    doc_type = data.get("doc_type")
    if doc_type not in _VALID_DOC_TYPES:
        return jsonify({"error": f"알 수 없는 문서 종류: {doc_type}"}), 400

    try:
        res = sb.table("documents").insert({
            "user_id": user_id,
            "doc_type": doc_type,
            "status": "draft",
            "current_step": 1,
        }).execute()
        row = (res.data or [{}])[0]
        return jsonify({
            "id": row.get("id"),
            "doc_type": doc_type,
            "status": "draft",
            "current_step": 1,
        }), 201
    except Exception as e:
        return jsonify({"error": f"문서 생성 실패: {e}"}), 500


@app.route("/api/documents/<doc_id>", methods=["PATCH"])
def api_documents_update(doc_id):
    """문서 업데이트 (변경 필드만). user_id 소유권 검증 (§3.5)."""
    user_id = _require_login()
    if not user_id:
        return jsonify({"error": "UNAUTHORIZED", "message": "로그인이 필요합니다."}), 401
    sb = _get_supabase()
    if sb is None:
        return jsonify({"error": "SUPABASE_UNAVAILABLE",
                        "message": "문서 저장소에 연결할 수 없습니다."}), 503

    data = request.get_json(silent=True) or {}
    patch = {k: v for k, v in data.items() if k in _PATCHABLE_FIELDS}
    if not patch:
        return jsonify({"error": "변경할 필드가 없습니다."}), 400
    if "status" in patch and patch["status"] not in _VALID_DOC_STATUS:
        return jsonify({"error": f"허용되지 않은 상태: {patch['status']}"}), 400

    try:
        res = (
            sb.table("documents").update(patch)
            .eq("id", doc_id).eq("user_id", user_id)
            .execute()
        )
        if not res.data:
            return jsonify({"error": "NOT_FOUND",
                            "message": "문서를 찾을 수 없거나 권한이 없습니다."}), 404
        return jsonify(res.data[0])
    except Exception as e:
        return jsonify({"error": f"문서 수정 실패: {e}"}), 500


@app.route("/api/documents/<doc_id>", methods=["DELETE"])
def api_documents_delete(doc_id):
    """문서 soft-delete (status → 'deleted'). user_id 소유권 검증 (§3.1)."""
    user_id = _require_login()
    if not user_id:
        return jsonify({"error": "UNAUTHORIZED", "message": "로그인이 필요합니다."}), 401
    sb = _get_supabase()
    if sb is None:
        return jsonify({"error": "SUPABASE_UNAVAILABLE",
                        "message": "문서 저장소에 연결할 수 없습니다."}), 503

    try:
        res = (
            sb.table("documents").update({"status": "deleted"})
            .eq("id", doc_id).eq("user_id", user_id)
            .execute()
        )
        if not res.data:
            return jsonify({"error": "NOT_FOUND",
                            "message": "문서를 찾을 수 없거나 권한이 없습니다."}), 404
        return jsonify({"success": True, "id": doc_id})
    except Exception as e:
        return jsonify({"error": f"문서 삭제 실패: {e}"}), 500


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5001, debug=True)
